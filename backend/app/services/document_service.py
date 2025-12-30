from io import BytesIO
from datetime import datetime, timedelta
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


def generate_proposal_docx(
    customer_name: str,
    customer_company: Optional[str],
    items: List[Dict],
    total: float,
    conditions: Optional[str],
    valid_days: int
) -> BytesIO:
    """
    Генерация коммерческого предложения в формате Word (DOCX).
    
    Args:
        customer_name: Имя клиента
        customer_company: Компания клиента
        items: Список товаров [{"name", "quantity", "price", "amount"}]
        total: Общая сумма
        conditions: Условия поставки
        valid_days: Срок действия КП в днях
    
    Returns:
        BytesIO buffer с DOCX файлом
    """
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Номер и дата
    doc.add_paragraph(f"№ КП-{datetime.now().strftime('%Y%m%d-%H%M')}")
    doc.add_paragraph(f"от {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph()
    
    # Клиент
    client_info = f"Клиент: {customer_name}"
    if customer_company:
        client_info += f" ({customer_company})"
    doc.add_paragraph(client_info)
    doc.add_paragraph()
    
    # Вступление
    doc.add_paragraph(
        "Благодарим за проявленный интерес к нашей продукции. "
        "Рады предложить вам следующие позиции:"
    )
    doc.add_paragraph()
    
    # Таблица товаров
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Заголовки
    header_cells = table.rows[0].cells
    headers = ['№', 'Наименование', 'Кол-во', 'Цена, ₽', 'Сумма, ₽']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Данные
    for idx, item in enumerate(items, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = item['name']
        row_cells[2].text = str(item['quantity'])
        row_cells[3].text = f"{item['price']:,.0f}"
        row_cells[4].text = f"{item['amount']:,.0f}"
    
    # Итого
    total_row = table.add_row().cells
    total_row[0].text = ""
    total_row[1].text = ""
    total_row[2].text = ""
    total_row[3].text = "ИТОГО:"
    total_row[3].paragraphs[0].runs[0].bold = True
    total_row[4].text = f"{total:,.0f}"
    total_row[4].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Условия
    if conditions:
        doc.add_paragraph(f"Условия: {conditions}")
    
    # Срок действия
    valid_until = datetime.now() + timedelta(days=valid_days)
    doc.add_paragraph(f"Предложение действительно до: {valid_until.strftime('%d.%m.%Y')}")
    
    doc.add_paragraph()
    doc.add_paragraph("С уважением,")
    doc.add_paragraph("Отдел продаж")
    
    # Сохраняем в buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def generate_proposal_pdf(
    customer_name: str,
    customer_company: Optional[str],
    items: List[Dict],
    total: float,
    conditions: Optional[str],
    valid_days: int
) -> BytesIO:
    """
    Генерация коммерческого предложения в формате PDF.
    
    Args:
        customer_name: Имя клиента
        customer_company: Компания клиента
        items: Список товаров [{"name", "quantity", "price", "amount"}]
        total: Общая сумма
        conditions: Условия поставки
        valid_days: Срок действия КП в днях
    
    Returns:
        BytesIO buffer с PDF файлом
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    
    styles = getSampleStyleSheet()
    
    # Создаём стили
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Title'],
        fontSize=18,
        spaceAfter=20,
        alignment=1  # Center
    )
    
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=10
    )
    
    elements = []
    
    # Заголовок
    elements.append(Paragraph("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ", title_style))
    elements.append(Spacer(1, 10))
    
    # Номер и дата
    elements.append(Paragraph(
        f"№ КП-{datetime.now().strftime('%Y%m%d-%H%M')} от {datetime.now().strftime('%d.%m.%Y')}",
        normal_style
    ))
    
    # Клиент
    client_info = f"<b>Клиент:</b> {customer_name}"
    if customer_company:
        client_info += f" ({customer_company})"
    elements.append(Paragraph(client_info, normal_style))
    elements.append(Spacer(1, 20))
    
    # Вступление
    elements.append(Paragraph(
        "Благодарим за проявленный интерес к нашей продукции. "
        "Рады предложить вам следующие позиции:",
        normal_style
    ))
    elements.append(Spacer(1, 15))
    
    # Таблица товаров
    table_data = [['№', 'Наименование', 'Кол-во', 'Цена, ₽', 'Сумма, ₽']]
    
    for idx, item in enumerate(items, 1):
        table_data.append([
            str(idx),
            item['name'],
            str(item['quantity']),
            f"{item['price']:,.0f}",
            f"{item['amount']:,.0f}"
        ])
    
    # Итого
    table_data.append(['', '', '', 'ИТОГО:', f"{total:,.0f}"])
    
    table = Table(table_data, colWidths=[1*cm, 8*cm, 2*cm, 3*cm, 3*cm])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -2), colors.beige),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('ALIGN', (2, 1), (-1, -1), 'RIGHT'),
        ('ALIGN', (1, 1), (1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (3, -1), (-1, -1), 'Helvetica-Bold'),
    ]))
    
    elements.append(table)
    elements.append(Spacer(1, 20))
    
    # Условия
    if conditions:
        elements.append(Paragraph(f"<b>Условия:</b> {conditions}", normal_style))
    
    # Срок действия
    valid_until = datetime.now() + timedelta(days=valid_days)
    elements.append(Paragraph(
        f"<b>Предложение действительно до:</b> {valid_until.strftime('%d.%m.%Y')}",
        normal_style
    ))
    
    elements.append(Spacer(1, 30))
    elements.append(Paragraph("С уважением,", normal_style))
    elements.append(Paragraph("<b>Отдел продаж</b>", normal_style))
    
    # Строим PDF
    doc.build(elements)
    buffer.seek(0)
    
    return buffer
